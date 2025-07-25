import os
import logging
from typing import List, Dict, Any, Optional, Tuple, TypedDict, Union
from dataclasses import dataclass
from pathlib import Path
from datetime import datetime
import json
import time
import streamlit as st
import re
from concurrent.futures import ThreadPoolExecutor
from document_chunk import DocumentChunk
from dateutil.parser import parse as date_parse

# Core libraries
from openai import AzureOpenAI
from openai import AzureOpenAI, OpenAI 
import fitz  # PyMuPDF
from docx import Document
import base64

# Vector store and embeddings
import numpy as np
from sentence_transformers import SentenceTransformer
import faiss
import pickle
from langchain.text_splitter import RecursiveCharacterTextSplitter

# LangGraph imports
from langgraph.graph import StateGraph, END
from langgraph.graph.message import add_messages
from langchain_core.messages import HumanMessage, AIMessage, BaseMessage

from typing_extensions import Annotated
from sentence_transformers.cross_encoder import CrossEncoder

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# File grouping configuration
FILE_GROUPS = {
    "Chemical Consumption": [
        "chemical", "consumption", "chemicals"
    ],
    "Engineering Tickets": [
        "engineering", "ticket", "tickets"
    ],
    "Risk Assessment and Hazard Analysis": [
        "risk", "hazard", "safety", "assessment", "analysis", "incident", "accident"
    ],
    "Well Recap": [
        "well recap", "Well Recap", "recap"
    ],
    "Mud Program": [
        "mud", "program"
    ],
    "Contractor Feedback": [
        "contractor", "feedback"
    ],
    "Hydraulic Summary": [
        "hydraulic", "summary","hydraulics"
    ],
    "Other Group": []  # Default group for unmatched files
}

@dataclass
class FileSystemInfo:
    """Stores file system structure information"""
    total_folders: int
    total_files: int
    file_types: Dict[str, int]
    folder_structure: Dict[str, Any]
    file_groups: Dict[str, List[str]]

# In your main script

class AgentState(TypedDict):
    messages: Annotated[List[BaseMessage], add_messages]
    original_question: str
    question: str
    selected_groups: List[str]
    search_results: List[Dict[str, Any]]
    final_answer: str
    confidence: float
    retrieval_grade: str
    correction_attempts: int
    critic_feedback: str
    refinement_attempts: int

    # Fields for date range filtering
    start_date: Optional[str]
    end_date: Optional[str]
    
    # Field for decomposed queries
    sub_queries: Optional[List[str]]

    # Field for structured entities
    query_entities: Optional[Dict[str, List[str]]]
    
    # Flag for fallback search
    fallback_triggered: Optional[bool]
    
    ## NEW: Field for the result of the direct file extraction path.
    direct_extraction_results: Optional[str]

class DocumentProcessor:
    """Enhanced document processor with hybrid layout-aware and OCR chunking for PDFs."""
    
    def __init__(self, azure_client: AzureOpenAI):
        """Initialize with Azure client and setup processing stats"""
        if not azure_client:
            raise ValueError("AzureOpenAI client cannot be None")
            
        self.azure_client = azure_client
        self.supported_extensions = {'.pdf', '.docx', '.txt'}
        self.processing_stats = {
            'text_extraction': [],
            'vision_ocr': [],
            'total_processed': 0
        }
        self.metadata_summary = []
    
    def classify_file_group(self, file_path: str) -> str:
        # -- UNCHANGED --
        file_name = os.path.basename(file_path).lower()
        for group_name, keywords in FILE_GROUPS.items():
            if group_name == "Other Group": continue
            for keyword in keywords:
                if keyword.lower() in file_name:
                    return group_name
        return "Other Group"

    def _format_table_as_markdown(self, table) -> str:
        # -- UNCHANGED --
        markdown_text = ""
        headers = [cell.strip() for cell in table.header.names if cell]
        if headers:
            markdown_text += "| " + " | ".join(headers) + " |\n"
            markdown_text += "| " + " | ".join(["---"] * len(headers)) + " |\n"
        for row in table.extract():
            cleaned_row = [str(cell).strip() if cell is not None else "" for cell in row]
            if any(cleaned_row):
                markdown_text += "| " + " | ".join(cleaned_row) + " |\n"
        return markdown_text

    ## NEW: Re-integrated the OCR method for processing scanned pages.
    def _extract_text_via_ocr_page(self, page) -> str:
        """Extract text from a single PDF page using GPT-4V OCR with rate limiting."""
        try:
            pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
            img_data = pix.tobytes("png")
            img_base64 = base64.b64encode(img_data).decode()
            
            max_retries = 3
            base_delay = 1.0
            for attempt in range(max_retries):
                try:
                    response = self.azure_client.chat.completions.create(
                        model="gpt-4o",
                        messages=[
                            {"role": "user", "content": [
                                {"type": "text", "text": "Extract all text from this image. Maintain the original formatting as much as possible. If there are tables, preserve their structure. Return only the extracted text without any commentary."},
                                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_base64}"}}
                            ]}
                        ],
                        max_tokens=4000, temperature=0
                    )
                    return response.choices[0].message.content
                except Exception as e:
                    if "429" in str(e):
                        delay = base_delay * (2 ** attempt)
                        logger.warning(f"Rate limited, retrying in {delay} seconds...")
                        time.sleep(delay)
                    else:
                        raise e
            logger.error(f"OCR failed after {max_retries} attempts.")
            return ""
        except Exception as e:
            logger.error(f"OCR failed for page: {e}")
            return ""

    ## MODIFIED: The process_pdf method is now a hybrid of layout-aware and OCR processing.
    def process_pdf(self, file_path: str, ocr_threshold: int = 50) -> List[DocumentChunk]:
        """
        MODIFIED: Process PDF using a hybrid approach.
        - For text-rich pages, it uses layout-aware chunking (tables, paragraphs).
        - For image-based pages (scans), it uses vision-based OCR.
        """
        if not os.path.exists(file_path) or os.path.getsize(file_path) == 0:
            return []

        all_chunks = []
        file_name = os.path.basename(file_path)
        file_group = self.classify_file_group(file_path)
        chunk_id_counter = 0

        try:
            doc = fitz.open(file_path)
            self.processing_stats['total_processed'] += 1
            
            for page_num, page in enumerate(doc):
                # HYBRID LOGIC: Decide whether to use layout-aware or OCR path
                if len(page.get_text("text").strip()) > ocr_threshold:
                    # --- PATH 1: Layout-Aware Chunking for Native PDFs ---
                    self.processing_stats['text_extraction'].append(page_num + 1)
                    
                    # Process tables
                    tables = page.find_tables()
                    table_bboxes = [fitz.Rect(t.bbox) for t in tables]
                    for i, table in enumerate(tables):
                        table_markdown = self._format_table_as_markdown(table)
                        if not table_markdown.strip(): continue
                        chunk_id = f"{file_name}_page_{page_num+1}_table_{i}"
                        metadata = {
                            "file_name": file_name, "group": file_group, "file_type": "pdf",
                            "element_type": "table", "page_number": page_num + 1,
                            "bounding_box": list(table.bbox), "processing_method": "Layout-Aware Extraction"
                        }
                        all_chunks.append(DocumentChunk(content=table_markdown, source_file=file_path, chunk_id=chunk_id, group=file_group, metadata=metadata))
                        chunk_id_counter += 1

                    # Process text blocks
                    text_blocks = page.get_text("blocks")
                    child_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=75)
                    for block in text_blocks:
                        block_rect = fitz.Rect(block[:4])
                        if any(block_rect.intersects(bbox) for bbox in table_bboxes): continue
                        block_text = block[4].strip()
                        if len(block_text) < 20: continue
                        
                        smaller_sub_chunks = child_splitter.split_text(block_text)
                        for sub_chunk_text in smaller_sub_chunks:
                            chunk_id = f"{file_name}_chunk_{chunk_id_counter}"
                            metadata = {
                                "parent_content": block_text, "file_name": file_name, "group": file_group,
                                "file_type": "pdf", "element_type": "paragraph", "page_number": page_num + 1,
                                "bounding_box": list(block[:4]), "processing_method": "Layout-Aware Extraction"
                            }
                            all_chunks.append(DocumentChunk(content=sub_chunk_text, source_file=file_path, chunk_id=chunk_id, group=file_group, metadata=metadata))
                            chunk_id_counter += 1
                else:
                    # --- PATH 2: OCR for Scanned Pages ---
                    logger.info(f"Page {page_num+1} in {file_name} has low text. Applying OCR.")
                    self.processing_stats['vision_ocr'].append(page_num + 1)
                    
                    ocr_text = self._extract_text_via_ocr_page(page)
                    if ocr_text:
                        # Use the standard _create_chunks for the OCR text stream
                        ocr_chunks = self._create_chunks(ocr_text, file_path, "pdf", file_group)
                        for chunk in ocr_chunks:
                            chunk.metadata['processing_method'] = "Vision (OCR)"
                            chunk.metadata['page_number'] = page_num + 1
                        all_chunks.extend(ocr_chunks)

            logger.info(f"Created {len(all_chunks)} hybrid chunks from {file_name}")
            return all_chunks

        except Exception as e:
            logger.error(f"Error in hybrid PDF processing for {file_path}: {e}", exc_info=True)
            return []

    def process_docx(self, file_path: str) -> List[DocumentChunk]:
        # -- UNCHANGED --
        try:
            doc = Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            for table in doc.tables:
                text += "\n" + "\n".join(" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows)
            return self._create_chunks(text, file_path, "docx", self.classify_file_group(file_path))
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            return []
    
    def process_txt(self, file_path: str) -> List[DocumentChunk]:
        # -- UNCHANGED --
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return self._create_chunks(file.read(), file_path, "txt", self.classify_file_group(file_path))
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}")
            return []

    def _extract_metadata(self, text: str) -> dict:
        """
        MODIFIED: Use dateutil.parser for robust, standardized date extraction.
        """
        # --- START OF FIX ---
        # The original regex is still useful for finding potential date strings
        date_pattern = r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b"
        potential_dates = re.findall(date_pattern, text)
        
        standardized_dates = set()
        for date_str in potential_dates:
            try:
                # Parse the date and format it to YYYY-MM-DD for reliable sorting and comparison
                dt = date_parse(date_str)
                standardized_dates.add(dt.strftime('%Y-%m-%d'))
            except (ValueError, OverflowError):
                # Ignore strings that look like dates but aren't (e.g., "1/2/34567")
                continue
        # --- END OF FIX ---

        well_pattern = r"\bWell\s*[A-Z0-9-]+\b"
        wells = list(set(re.findall(well_pattern, text, re.IGNORECASE)))
        
        # Return the standardized, sorted list of dates
        return {"dates": sorted(list(standardized_dates)), "wells": wells}

    def _create_chunks(self, text: str, file_path: str, file_type: str, group: str) -> List[DocumentChunk]:
        # -- UNCHANGED --
        if not text.strip(): return []
        file_name = os.path.basename(file_path)
        metadata = self._extract_metadata(text)
        parent_chunks = text.split("\n\n")
        child_chunks = []
        chunk_id_counter = 0
        child_splitter = RecursiveCharacterTextSplitter(chunk_size=400, chunk_overlap=50)
        for i, parent_chunk_text in enumerate(parent_chunks):
            if not parent_chunk_text.strip(): continue
            smaller_sub_chunks = child_splitter.split_text(parent_chunk_text)
            for sub_chunk_text in smaller_sub_chunks:
                chunk_id = f"{file_name}_chunk_{chunk_id_counter}"
                new_chunk = DocumentChunk(
                    content=sub_chunk_text, source_file=file_path, chunk_id=chunk_id, group=group,
                    metadata={
                        "parent_content": parent_chunk_text, "parent_chunk_index": i,
                        "file_type": file_type, "file_size": len(sub_chunk_text),
                        "processed_at": datetime.now().isoformat(), "file_name": file_name,
                        "group": group, "processing_method": "Text Extraction",
                        "dates": metadata["dates"], "wells": metadata["wells"]
                    }
                )
                child_chunks.append(new_chunk)
                chunk_id_counter += 1
        return child_chunks

class GroupedVectorStore:
    """Manages document embeddings with group-based organization"""
    
    def __init__(self, embedding_model_name: str = "all-mpnet-base-v2"):
        """Initialize with configurable embedding model"""
        # Store the model name for easy access later
        self.model_name = embedding_model_name
        self.embedding_model = SentenceTransformer(self.model_name)
        self.group_indices = {}
        self.group_chunks = {}
        self.dimension = self.embedding_model.get_sentence_embedding_dimension()
        
        # Enhanced tracking
        self.file_type_stats = {}
        self.group_file_details = {}
        self.date_index = {}  # {date: [file1, file2]}
        self.well_index = {}  # {well_name: [file1, file2]}
        
        logger.info(f"Initialized GroupedVectorStore with model: {self.model_name}")
        
    def _update_indices(self, chunk: DocumentChunk):
        """Update date and well indices"""
        file_name = chunk.metadata.get('file_name', 'unknown')
        dates = chunk.metadata.get('dates', [])
        wells = chunk.metadata.get('wells', [])
        
        for date in dates:
            if date not in self.date_index:
                self.date_index[date] = []
            if file_name not in self.date_index[date]:
                self.date_index[date].append(file_name)
                
        for well in wells:
            if well not in self.well_index:
                self.well_index[well] = []
            if file_name not in self.well_index[well]:
                self.well_index[well].append(file_name)
    
    def add_documents(self, chunks: List[DocumentChunk]):
        """MODIFIED: Enhanced document addition to correctly track multiple processing methods per file."""
        if not chunks:
            logger.warning("No chunks provided to add_documents")
            return
        
        for chunk in chunks:
            file_type = chunk.metadata.get('file_type', 'unknown')
            group = chunk.group
            file_name = chunk.metadata.get('file_name', 'unknown')
            
            # This logic remains the same
            if file_type not in self.file_type_stats:
                self.file_type_stats[file_type] = 0
            self.file_type_stats[file_type] += 1
            
            if group not in self.group_file_details:
                self.group_file_details[group] = {
                    'files': set(), 'file_details': {}, 'total_chunks': 0
                }
            
            self.group_file_details[group]['files'].add(file_name)
            self.group_file_details[group]['total_chunks'] += 1
            
            # --- START OF FIX ---
            # Correctly handle multiple processing methods for a single file
            
            # Initialize the file details entry if it's the first time we see this file
            if file_name not in self.group_file_details[group]['file_details']:
                self.group_file_details[group]['file_details'][file_name] = {
                    'file_type': file_type,
                    'processing_methods': set(), # Use a set to store unique methods
                    'chunks': 0,
                    'dates': set(),
                    'wells': set()
                }
            
            # Add the processing method from the current chunk to the set for this file
            current_method = chunk.metadata.get('processing_method', 'Unknown')
            self.group_file_details[group]['file_details'][file_name]['processing_methods'].add(current_method)
            
            # Update chunk count and other metadata
            self.group_file_details[group]['file_details'][file_name]['chunks'] += 1
            self.group_file_details[group]['file_details'][file_name]['dates'].update(chunk.metadata.get('dates', []))
            self.group_file_details[group]['file_details'][file_name]['wells'].update(chunk.metadata.get('wells', []))
            
            # --- END OF FIX ---
            
            self._update_indices(chunk)
            
        # Grouping and embedding logic remains the same
        grouped_chunks = {}
        for chunk in chunks:
            group = chunk.group
            if group not in grouped_chunks:
                grouped_chunks[group] = []
            grouped_chunks[group].append(chunk)
            
        for group_name, group_chunks in grouped_chunks.items():
            logger.info(f"Adding {len(group_chunks)} chunks to group: {group_name}")
            self._add_to_group(group_name, group_chunks)
    
    def _add_to_group(self, group_name: str, chunks: List[DocumentChunk]): 

        """Add chunks to a group's vector store with efficient batch embeddings""" 

        if not chunks: 

            return 

    

        # Initialize index if needed 

        if group_name not in self.group_indices: 

            self.group_indices[group_name] = faiss.IndexFlatIP(self.dimension) 

            self.group_chunks[group_name] = [] 

    

        # Process in batches using native batch processing 

        batch_size = 32  # Optimal batch size for sentence transformers 

        embeddings = [] 

        

        for i in range(0, len(chunks), batch_size): 

            batch = chunks[i:i + batch_size] 

            try: 

                # Extract content for batch processing 

                batch_contents = [chunk.content for chunk in batch] 

                

                # Use native batch encoding (much more efficient than ThreadPoolExecutor) 

                batch_embeddings = self.embedding_model.encode( 

                    batch_contents, 

                    batch_size=batch_size, 

                    show_progress_bar=False, 

                    convert_to_numpy=True 

                ) 

                

                # Ensure embeddings are in the correct format 

                if batch_embeddings.ndim == 1: 

                    batch_embeddings = batch_embeddings.reshape(1, -1) 

                

                embeddings.extend(batch_embeddings) 

                

                # Store embeddings in chunks 

                for chunk, emb in zip(batch, batch_embeddings): 

                    chunk.embedding = emb 

                    

            except Exception as e: 

                logger.error(f"Embedding generation failed for batch {i//batch_size}: {e}") 

                # Continue processing other batches 

                continue 

    

        if embeddings: 

            try: 

                # Convert to numpy array and ensure correct dtype 

                embeddings_array = np.array(embeddings, dtype='float32') 

                

                # Add to FAISS index 

                self.group_indices[group_name].add(embeddings_array) 

                self.group_chunks[group_name].extend(chunks) 

                

                logger.info(f"Added {len(embeddings)} embeddings to group {group_name}") 

                

            except Exception as e: 

                logger.error(f"Failed to add embeddings to FAISS index for group {group_name}: {e}") 
    
    def search_groups(self, query: str, groups: List[str], k: int = 5) -> List[Tuple[DocumentChunk, float, str]]:
        """Search across specified groups"""
        all_results = []
        
        for group_name in groups:
            # Skip if group doesn't exist or is empty
            if group_name not in self.group_indices or group_name not in self.group_chunks or len(self.group_chunks.get(group_name, [])) == 0:
                logger.warning(f"Group {group_name} is empty or not initialized")
                continue
            
            try:
                query_embedding = self.embedding_model.encode(query)
                scores, indices = self.group_indices[group_name].search(
                    np.array([query_embedding]).astype('float32'), 
                    min(k, len(self.group_chunks[group_name]))
                )
                
                for score, idx in zip(scores[0], indices[0]):
                    if 0 <= idx < len(self.group_chunks[group_name]):
                        all_results.append((
                            self.group_chunks[group_name][idx], 
                            float(score), 
                            group_name
                        ))
                        
            except Exception as e:
                logger.error(f"Error searching group {group_name}: {e}")
                continue
        
        # Sort by score and return top k
        all_results.sort(key=lambda x: x[1], reverse=True)
        logger.info(f"Found {len(all_results)} results across groups {groups}")
        return all_results[:k]
    
    def get_enhanced_group_stats(self) -> Dict[str, Any]:
        """Get enhanced statistics with file details"""
        stats = {
            'file_type_distribution': dict(self.file_type_stats),
            'total_groups': len(self.group_indices),
            'groups': {}
        }
        
        for group_name in self.group_indices:
            group_details = self.group_file_details.get(group_name, {})
            chunk_count = len(self.group_chunks.get(group_name, []))
            
            stats['groups'][group_name] = {
                'chunk_count': chunk_count,
                'unique_files': len(group_details.get('files', set())),
                'file_list': list(group_details.get('files', set())),
                'file_details': group_details.get('file_details', {}),
                'total_chunks_tracked': group_details.get('total_chunks', 0)
            }
        
        return stats

    def get_file_type_summary(self) -> Dict[str, Any]:
        """Get summary by file type"""
        summary = {
            'pdf_files': 0,
            'docx_files': 0,
            'txt_files': 0,
            'total_files': 0,
            'processing_methods': {
                'text_extraction': 0,
                'vision_ocr': 0
            }
        }
        
        all_processed_files = set()

        # --- FIX WAS APPLIED HERE ---
        # The '.vector_store' part was removed from the line below
        for group_name, group_details in self.group_file_details.items():
            for file_name, file_info in group_details.get('file_details', {}).items():
                
                # This check is now redundant since you are iterating through all_processed_files later
                # if file_name in all_processed_files:
                #    continue
                # all_processed_files.add(file_name)
                
                file_type = file_info.get('file_type', 'unknown')
                
                if file_type == 'pdf':
                    summary['pdf_files'] += 1
                elif file_type == 'docx':
                    summary['docx_files'] += 1
                elif file_type == 'txt':
                    summary['txt_files'] += 1
                
                summary['total_files'] += 1
                
                # This part correctly checks the set of processing methods for a file
                processing_methods_used = file_info.get('processing_methods', set())
                if 'Layout-Aware Extraction' in processing_methods_used:
                    summary['processing_methods']['text_extraction'] += 1
                if 'Vision (OCR)' in processing_methods_used:
                    summary['processing_methods']['vision_ocr'] += 1
        
        return summary
    

    def get_group_files(self, group_name: str) -> Dict[str, Any]:
        """Get detailed information about files in a specific group"""
        if group_name not in self.group_file_details:
            return {'files': [], 'total_files': 0}
        
        group_info = self.group_file_details[group_name]
        return {
            'files': list(group_info.get('files', set())),
            'total_files': len(group_info.get('files', set())),
            'file_details': group_info.get('file_details', {}),
            'total_chunks': group_info.get('total_chunks', 0)
        }
    
    def save_group(self, group_name: str, folder_path: str):
        """Save a specific group's vector store to a folder, skipping empty groups."""
        # --- START OF FIX ---
        # Check if the group exists and has content before attempting to save.
        if group_name not in self.group_chunks or not self.group_chunks.get(group_name):
            logger.info(f"Skipping save for empty group: {group_name}")
            return
        # --- END OF FIX ---

        try:
            os.makedirs(folder_path, exist_ok=True)
            
            if group_name in self.group_indices:
                index_path = os.path.join(folder_path, f"{group_name}.index")
                faiss.write_index(self.group_indices[group_name], index_path)
                logger.info(f"Saved FAISS index for group {group_name} to {index_path}")
            
            if group_name in self.group_chunks:
                chunks_to_save = []
                for chunk in self.group_chunks[group_name]:
                    chunk_copy = DocumentChunk(
                        content=chunk.content,
                        source_file=chunk.source_file,
                        chunk_id=chunk.chunk_id,
                        group=chunk.group,
                        metadata=chunk.metadata,
                        embedding=None
                    )
                    chunks_to_save.append(chunk_copy)
                
                chunks_path = os.path.join(folder_path, f"{group_name}.chunks")
                with open(chunks_path, 'wb') as f:
                    pickle.dump(chunks_to_save, f)
                
                logger.info(f"Saved {len(chunks_to_save)} chunks for group {group_name} to {chunks_path}")
                
        except Exception as e:
            logger.error(f"Error saving group {group_name}: {e}")
    
    def save_all_groups(self, folder_path: str):
        """Save all groups to a single folder"""
        
        try:
            os.makedirs(folder_path, exist_ok=True)
            for group_name in self.group_indices:
                self.save_group(group_name, folder_path)
            logger.info(f"Saved all groups to folder: {folder_path}")
            return True
        except Exception as e:
            logger.error(f"Error saving all groups: {e}")
            return False
    
    
    def load_group(self, group_name: str, folder_path: str):
        """Load a specific group's vector store from folder"""
        try:
            index_path = os.path.join(folder_path, f"{group_name}.index")
            chunks_path = os.path.join(folder_path, f"{group_name}.chunks")
            
            if os.path.exists(index_path):
                self.group_indices[group_name] = faiss.read_index(index_path)
                logger.info(f"Loaded FAISS index for group {group_name} from {index_path}")
            else:
                logger.warning(f"FAISS index file not found at {index_path}")
                return False
            
            if os.path.exists(chunks_path):
                with open(chunks_path, 'rb') as f:
                    self.group_chunks[group_name] = pickle.load(f)
                logger.info(f"Loaded {len(self.group_chunks[group_name])} chunks for group {group_name} from {chunks_path}")
            else:
                logger.warning(f"Chunks file not found at {chunks_path}")
                return False
            
            return True
            
        except Exception as e:
            logger.error(f"Error loading group {group_name}: {e}")
            return False
    
    # In class GroupedVectorStore:

    def load_all_groups(self, folder_path: str):
        """
        MODIFIED: Load all available groups from folder and rebuild metadata indices.
        """
        loaded_groups = []
        
        if not os.path.exists(folder_path):
            logger.error(f"Vector store folder not found: {folder_path}")
            return loaded_groups
            
        for filename in os.listdir(folder_path):
            if filename.endswith('.index'):
                group_name = filename[:-6]  # Remove '.index'
                if self.load_group(group_name, folder_path):
                    loaded_groups.append(group_name)

        # --- START OF FIX: Rebuild metadata after loading ---
        if loaded_groups:
            logger.info("Rebuilding metadata from loaded chunks...")
            all_loaded_chunks = []
            for group_name in loaded_groups:
                all_loaded_chunks.extend(self.group_chunks.get(group_name, []))
            
            # Use a simplified version of the add_documents logic to rebuild stats
            for chunk in all_loaded_chunks:
                file_type = chunk.metadata.get('file_type', 'unknown')
                group = chunk.group
                file_name = chunk.metadata.get('file_name', 'unknown')

                # Update file type stats
                self.file_type_stats[file_type] = self.file_type_stats.get(file_type, 0) + 1
                
                # Initialize group details if not present
                if group not in self.group_file_details:
                    self.group_file_details[group] = {
                        'files': set(), 'file_details': {}, 'total_chunks': 0
                    }
                
                self.group_file_details[group]['files'].add(file_name)
                self.group_file_details[group]['total_chunks'] += 1
                
                if file_name not in self.group_file_details[group]['file_details']:
                     self.group_file_details[group]['file_details'][file_name] = {
                        'file_type': file_type,
                        'processing_method': chunk.metadata.get('processing_method', 'Unknown'),
                        'chunks': 0
                    }
                self.group_file_details[group]['file_details'][file_name]['chunks'] += 1
                
                # Rebuild date and well indices
                self._update_indices(chunk)
            logger.info("Metadata rebuilt successfully.")
        # --- END OF FIX ---

        logger.info(f"Loaded groups from {folder_path}: {loaded_groups}")
        return loaded_groups


class FileSystemAnalyzer:
    """Analyzes file system structure with group classification"""
    
    def __init__(self, root_path: str):
        self.root_path = Path(root_path)
        self.file_info = None
        self.document_processor = None
    
    def set_document_processor(self, processor: DocumentProcessor):
        """Set document processor for file classification"""
        self.document_processor = processor
    
    def analyze(self) -> FileSystemInfo:
        """Analyze file system structure with group classification"""
        folder_count = 0
        file_count = 0
        file_types = {}
        folder_structure = {}
        file_groups = {group: [] for group in FILE_GROUPS.keys()}
        
        try:
            for root, dirs, files in os.walk(self.root_path):
                folder_count += len(dirs)
                file_count += len(files)
                
                for file in files:
                    file_path = os.path.join(root, file)
                    ext = Path(file).suffix.lower()
                    file_types[ext] = file_types.get(ext, 0) + 1
                    
                    # Classify file into group
                    if self.document_processor:
                        group = self.document_processor.classify_file_group(file_path)
                        file_groups[group].append(file)
                
                rel_path = os.path.relpath(root, self.root_path)
                if rel_path == '.':
                    rel_path = 'root'
                
                folder_structure[rel_path] = {
                    'subdirs': dirs,
                    'files': files,
                    'file_count': len(files)
                }
        
        except Exception as e:
            logger.error(f"Error analyzing file system: {e}")
        
        self.file_info = FileSystemInfo(
            total_folders=folder_count,
            total_files=file_count,
            file_types=file_types,
            folder_structure=folder_structure,
            file_groups=file_groups
        )
        
        logger.info(f"Analyzed file system: {file_count} files, {folder_count} folders")
        return self.file_info
    
    def get_folder_summary(self) -> str:
        """Get a summary of the folder structure with group information"""
        if not self.file_info:
            self.analyze()
        
        summary = f"File System Analysis:\n"
        summary += f"- Total folders: {self.file_info.total_folders}\n"
        summary += f"- Total files: {self.file_info.total_files}\n"
        summary += f"- File types: {dict(self.file_info.file_types)}\n\n"
        
        summary += "File Groups:\n"
        for group, files in self.file_info.file_groups.items():
            summary += f"- {group}: {len(files)} files\n"
        
        return summary

## MODIFIED: This class is heavily refactored to implement the advanced agentic loop.
class LangGraphAgent:
    """LangGraph-based agent with self-correcting retrieval and proactive synthesis."""
    
    ## MODIFIED: __init__ now accepts and stores the Azure deployment name.
    def __init__(self, azure_client: AzureOpenAI, deepseek_client: OpenAI, vector_store: GroupedVectorStore, file_analyzer: FileSystemAnalyzer, azure_deployment_name: str):
        self.azure_client = azure_client
        self.deepseek_client = deepseek_client
        self.vector_store = vector_store
        self.file_analyzer = file_analyzer
        self.cross_encoder = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')
        self.query_cache = {}
        
        # ## NEW: Store the deployment name for use in all Azure OpenAI calls.
        self.azure_deployment_name = azure_deployment_name 
        
        self.graph = self._build_graph()

    def _format_history(self, messages: List[BaseMessage]) -> str:
        """Formats the message history into a string for prompts."""
        if not messages:
            return "No conversation history."
        # Exclude the most recent message since it's the current question
        history = messages[:-1]
        if not history:
            return "No previous conversation."

        return "\n".join([f"{'User' if isinstance(m, HumanMessage) else 'Assistant'}: {m.content}" for m in history])

    def _criticize_answer(self, state: AgentState) -> AgentState:
        """Uses the live DeepSeek API to criticize the generated answer."""
        logger.info("Criticizing the generated answer with DeepSeek API...")
        
        question = state["original_question"]
        current_answer = state["final_answer"]
        
        if not state.get("search_results"):
            state["critic_feedback"] = "No improvement needed."
            return state
            
        context_str = "\n---\n".join([res['content_for_llm'] for res in state["search_results"]])
        
        critic_prompt = f"""You are a strict, helpful, and precise quality-control assistant. Your task is to evaluate a generated answer based on a given context and question.

        Context from retrieved documents:
        ---
        {context_str}
        ---

        Original Question: "{question}"

        Generated Answer to Evaluate:
        "{current_answer}"

        Critique Guidelines:
        1.  **Faithfulness:** Is the answer fully supported by the provided context? Point out any claims that are not backed by the documents.
        2.  **Completeness:** Does the answer fully address the user's question? If it's missing key details that are present in the context, mention them.
        3.  **Conciseness:** Is the answer direct or does it contain unnecessary fluff?
        4.  **Actionable Feedback:** Provide clear, actionable instructions on how to improve the answer.

        If the answer is perfect and requires no changes, respond with ONLY the phrase "No improvement needed."
        Otherwise, provide your critique.

        Critique:"""
        
        try:
            response = self.deepseek_client.chat.completions.create(
                model="deepseek-chat",
                messages=[{"role": "user", "content": critic_prompt}],
                max_tokens=500,
                temperature=0.0
            )
            feedback = response.choices[0].message.content
        except Exception as e:
            logger.error(f"DeepSeek API call failed: {e}")
            feedback = "No improvement needed." 
            
        state["critic_feedback"] = feedback
        state["refinement_attempts"] = state.get("refinement_attempts", 0) + 1
        logger.info(f"Critic Feedback from DeepSeek: {feedback}")
        
        return state

    def _identify_and_extract_files(self, state: AgentState) -> dict:
        """
        Identifies the most relevant files for a query and extracts their full content.
        This provides broader context to complement the targeted vector search.
        """
        logger.info("Starting direct file identification and extraction...")
        question = state["original_question"]
        selected_groups = state["selected_groups"]

        # 1. Get a list of all available files in the selected groups.
        available_files = []
        for group_name in selected_groups:
            group_details = self.vector_store.get_group_files(group_name)
            available_files.extend(group_details.get('files', []))
        
        if not available_files:
            logger.warning("No files available in selected groups for direct extraction.")
            return {"direct_extraction_results": None}

        # 2. Use an LLM to identify the most relevant files.
        identification_prompt = f"""
        You are a file system expert. Based on the user's question, identify the single most relevant file from the list below that is most likely to contain the answer.

        User Question: "{question}"

        List of Available Files:
        - {chr(10).join(available_files)}

        Instructions:
        Return a single JSON object with one key, "best_file", containing the full name of the most relevant file.
        
        JSON Output:
        """
        
        try:
            response = self.azure_client.chat.completions.create(
                model=self.azure_deployment_name,
                messages=[{"role": "user", "content": identification_prompt}],
                max_tokens=200,
                temperature=0,
                response_format={"type": "json_object"}
            )
            result = json.loads(response.choices[0].message.content)
            best_file = result.get("best_file")

            if not best_file or best_file not in available_files:
                logger.warning("LLM failed to identify a valid file. Skipping direct extraction.")
                return {"direct_extraction_results": None}
                
            logger.info(f"LLM identified '{best_file}' as the most relevant.")

            # 3. "Re-read" the file by concatenating all its chunks.
            file_chunks = []
            for group_name in selected_groups:
                if group_name in self.vector_store.group_chunks:
                    for chunk in self.vector_store.group_chunks[group_name]:
                        if chunk.metadata.get('file_name') == best_file:
                            file_chunks.append(chunk.content)
            
            if not file_chunks:
                logger.warning(f"Could not find any chunks for the identified file: {best_file}")
                return {"direct_extraction_results": None}

            full_content = "\n\n---\n\n".join(file_chunks)
            extraction_summary = f"## Context from Full Document: {best_file}\n\n{full_content}"
            
            logger.info(f"Successfully extracted full content from {best_file}.")
            return {"direct_extraction_results": extraction_summary}

        except Exception as e:
            logger.error(f"Error during direct file extraction: {e}", exc_info=True)
            return {"direct_extraction_results": None}

    def _query_transformer(self, state: AgentState) -> dict:
        state['correction_attempts'] = state.get('correction_attempts', 0) + 1
        logger.info(f"Transforming query into investigative sub-queries (Attempt {state['correction_attempts']})")

        group_descriptions = "\n".join([f"- **{name}:** {', '.join(keywords)}" for name, keywords in FILE_GROUPS.items() if keywords])
        question = state['original_question']
        conversation_history = self._format_history(state["messages"])
        
        prompt = f"""You are an expert research analyst. Your task is to decompose a user's high-level question into a strategic search plan, considering the conversation history for context.

        Available Document Groups:
        {group_descriptions}

        Conversation History:
        ---
        {conversation_history}
        ---

        Current User Question: "{question}"

        Instructions:
        1.  **Analyze Context:** Use the conversation history to understand the full context of the user's current question.
        2.  **Formulate Investigative Queries:** Create precise queries to find evidence for the current question. The queries should use keywords likely to be found in the target documents.
        
        Return your final output as a single JSON object with a single key "queries" containing a list of the search strings.

        JSON Output:
        """
        
        response = self.azure_client.chat.completions.create(
            model=self.azure_deployment_name,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=500,
            temperature=0.0,
            response_format={"type": "json_object"}
        )
        
        try:
            result = json.loads(response.choices[0].message.content)
            sub_queries = result.get("queries", [state["original_question"]])
            logger.info(f"Decomposed query into: {sub_queries}")
            return {"sub_queries": sub_queries, "correction_attempts": state['correction_attempts']}
        except (json.JSONDecodeError, KeyError) as e:
            logger.error(f"Failed to decode sub-queries: {e}. Falling back to original question.")
            return {"sub_queries": [state["original_question"]], "correction_attempts": state['correction_attempts']}

    ## NEW: Node for grading the relevance of retrieved documents.
    def _grade_documents(self, state: AgentState) -> AgentState:
        # This method remains unchanged.
        logger.info("Grading retrieved documents...")
        question = state["original_question"]
        documents = state["search_results"]
        
        if not documents:
            state["retrieval_grade"] = "not_relevant"
            logger.warning("No documents found to grade.")
            return state

        avg_score = sum(doc['score'] for doc in documents) / len(documents)
        
        if avg_score > 0.5: # This threshold can be tuned
            logger.info("Grading: Documents are RELEVANT.")
            state["retrieval_grade"] = "relevant"
        else:
            logger.warning("Grading: Documents are NOT RELEVANT. Average score was low.")
            state["retrieval_grade"] = "not_relevant"
            
        return state

    ## NEW: Conditional edge logic for the self-correction loop.
    def _decide_to_continue(self, state: AgentState) -> str:
        # This method remains unchanged.
        attempts = state.get("correction_attempts", 0)
        max_attempts = 2

        if not state["search_results"] and state.get("start_date"):
            logger.warning("Decision: Date filter removed all documents. Forcing synthesis to inform the user.")
            return "synthesize"

        if state["retrieval_grade"] == "relevant" or attempts >= max_attempts:
            if state["retrieval_grade"] != "relevant":
                logger.warning(f"Decision: Max correction attempts ({max_attempts}) reached. Forcing synthesis.")
            else:
                logger.info("Decision: Retrieval successful. Proceeding to synthesis.")
            return "synthesize"
        else:
            logger.info(f"Decision: Retrieval failed. Attempting to re-transform query (Attempt {attempts + 1}).")
            return "transform_query"

    ## MODIFIED: Synthesis prompt now asks for quote-level citations.
    def _synthesize_answer(self, state: AgentState) -> AgentState:
        logger.info("Synthesizing final answer using dual-path context...")
        question = state["original_question"]
        search_results = state["search_results"]
        conversation_history = self._format_history(state["messages"])
        
        ## NEW: Get the results from the direct extraction path.
        direct_extraction = state.get("direct_extraction_results")

        if not search_results and not direct_extraction:
            state["final_answer"] = "I could not find any relevant information in the documents to answer your question."
            state["confidence"] = 0.0
            return state

        # Build context from vector search (Path A)
        vector_context_parts = []
        for result in search_results:
            context_for_llm = result['content_for_llm']
            file_name = result['file_name']
            chunk_id = result['chunk_id']
            vector_context_parts.append(f"Source (File: {file_name}, Chunk ID: {chunk_id}):\n{context_for_llm}\n---")
        vector_context_str = "\n".join(vector_context_parts)
        
        # Combine the contexts
        combined_context = ""
        if vector_context_str:
            combined_context += "### Context from Precise Vector Search\n" + vector_context_str
        
        if direct_extraction:
            combined_context += "\n\n" + direct_extraction # This string already has its own header

        synthesis_prompt = f"""You are an expert Q&A system that synthesizes information from two sources to provide a single, highly reliable answer.

        **Task:**
        Answer the user's question by combining information from the "Precise Vector Search" (which contains specific, relevant snippets) and the "Context from Full Document" (which provides broader context). Use ONLY the provided materials. Your entire response must follow the strict format below.

        **Conversation History:**
        {conversation_history}

        **Current User Question:** "{question}"

        **Retrieved Information:**
        ---
        {combined_context}
        ---

        **Output Format (Strict):**

        [Start with a direct, one-sentence answer to the question, with the key finding in bold. Synthesize the details from both the precise snippets and the full document context to create a comprehensive explanation. Cite your sources in this section using <cite id="chunk_id_goes_here"></cite> for snippets from the vector search.]
        ---

        ### Citation
        * **File Name**: `[List the name of the file or files the primary information came from, e.g., Well Summary UZ-034.pdf]`

        ### More Information
        * **Supporting Quote**: `[Quote the specific text from the document that directly supports the main answer. Prefer quotes from the Precise Vector Search context if available.]`

        **Your Response:**
        """

        try:
            response = self.azure_client.chat.completions.create(
                model=self.azure_deployment_name,
                messages=[{"role": "user", "content": synthesis_prompt}],
                max_tokens=1500, temperature=0.0
            )
            answer = response.choices[0].message.content
            
            avg_score = sum(r.get("score", 0) for r in search_results) / len(search_results) if search_results else 0
            confidence = 1 / (1 + np.exp(-avg_score))
            
            state["final_answer"] = answer
            state["confidence"] = confidence
            logger.info(f"Generated dual-path answer with confidence: {confidence:.2f}")
            return state
        except Exception as e:
            logger.error(f"Error in answer synthesis: {e}")
            state["final_answer"] = f"I encountered an error while generating the answer: {str(e)}"
            state["confidence"] = 0.0
            return state

    # In class LangGraphAgent

    def _build_graph(self) -> StateGraph:
        workflow = StateGraph(AgentState)

        workflow.add_node("group_selector", self._select_groups)
        workflow.add_node("extract_date_range", self._extract_query_date_range)
        workflow.add_node("extract_entities", self._extract_entities)
        workflow.add_node("transform_query", self._query_transformer)
        workflow.add_node("document_searcher", self._search_documents)
        workflow.add_node("grade_documents", self._grade_documents)
        
        ## NEW: Add the direct extraction node to the graph.
        workflow.add_node("identify_and_extract", self._identify_and_extract_files)
        
        workflow.add_node("answer_synthesizer", self._synthesize_answer)
        workflow.add_node("criticize_answer", self._criticize_answer)
        workflow.add_node("refine_answer", self._refine_answer)

        workflow.set_entry_point("group_selector")
        workflow.add_edge("group_selector", "extract_date_range")
        workflow.add_edge("extract_date_range", "extract_entities")
        workflow.add_edge("extract_entities", "transform_query")
        workflow.add_edge("transform_query", "document_searcher")
        workflow.add_edge("document_searcher", "grade_documents")

        workflow.add_conditional_edges(
            "grade_documents", self._decide_to_continue,
            {
                ## MODIFIED: If synthesis is chosen, go to direct extraction first.
                "synthesize": "identify_and_extract", 
                "transform_query": "transform_query"
            }
        )
        
        ## MODIFIED: The synthesizer now runs after the direct extraction is complete.
        workflow.add_edge("identify_and_extract", "answer_synthesizer")

        workflow.add_edge("answer_synthesizer", "criticize_answer")

        def _decide_to_refine(state: AgentState) -> str:
            feedback = state.get("critic_feedback", "")
            attempts = state.get("refinement_attempts", 0)
            if "no improvement needed" in feedback.lower() or attempts >= 2:
                logger.info("Decision: Answer is good or max refinement attempts reached. Finishing.")
                return "finish"
            else:
                logger.info("Decision: Answer requires refinement.")
                return "refine"

        workflow.add_conditional_edges(
            "criticize_answer",
            _decide_to_refine,
            {"refine": "refine_answer", "finish": END}
        )

        workflow.add_edge("refine_answer", "criticize_answer")

        compiled_graph = workflow.compile(checkpointer=None)
        compiled_graph.recursion_limit = 20
        return compiled_graph

    def _answer_structure_query(self, question: str) -> Dict[str, Any]:
        # This method remains unchanged.
        logger.info(f"Answering structure query: {question}")
        if not self.file_analyzer or not self.file_analyzer.file_info:
            return {"answer": "The file system has not been analyzed yet. Please ingest a folder first.", "confidence": 1.0, "search_results": []}

        context = self.file_analyzer.get_folder_summary()

        prompt = f"""You are a helpful assistant that answers questions about a file system's structure.
        Use ONLY the provided context to answer the user's question. Do not make up information.

        Context:
        ---
        {context}
        ---

        User Question: "{question}"

        Answer:"""

        try:
            response = self.azure_client.chat.completions.create(
                model=self.azure_deployment_name,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=500,
                temperature=0
            )
            answer = response.choices[0].message.content
            return {
                "answer": answer,
                "confidence": 0.95,
                "selected_groups": [],
                "search_results": [],
            }
        except Exception as e:
            logger.error(f"Error answering structure query: {e}")
            return {"answer": f"I encountered an error while answering the structure query: {e}", "confidence": 0.0, "search_results": []}

    ## MODIFIED: This is the main entry point to the agent. It now takes the full message history.
    def process_query(self, messages: List[BaseMessage]) -> Dict[str, Any]:
        """Main query processor with the new agentic workflow and chat history."""
        if not messages or not isinstance(messages[-1], HumanMessage):
            return {
                "answer": "Please ask a question to begin.", "confidence": 0.0,
                "selected_groups": [], "search_results": []
            }
        
        # The current question is the last message. The history is the full list.
        question = messages[-1].content
        
        # Simple cache key based on the last question. More complex caching could hash the history.
        cache_key = question.lower().strip()
        if cache_key in self.query_cache:
            logger.info(f"Returning cached result for query: {question}")
            return self.query_cache[cache_key]

        # <-- Router to decide the query type -->
        routing_prompt = f"""Is the following user question about the file system structure (like file counts, folder names, file types, or file groups) or about the text content inside the documents?
        Respond with a single word: 'structure' or 'content'.

        Question: "{question}"
        """
        try:
            response = self.azure_client.chat.completions.create(
                model=self.azure_deployment_name,
                messages=[{"role": "system", "content": routing_prompt}],
                max_tokens=5,
                temperature=0
            )
            route = response.choices[0].message.content.lower().strip()
        except Exception as e:
            logger.error(f"Routing failed: {e}. Defaulting to content search.")
            route = "content"

        if "structure" in route:
            result = self._answer_structure_query(question)
            self.query_cache[cache_key] = result
            return result
        
        # --- Run the LangGraph workflow ---
        initial_state = {
            "messages": messages,
            "original_question": question,
            "question": question,
            "correction_attempts": 0,
            "refinement_attempts": 0, # Reset refinement attempts for each new query
        }
        
        try:
            final_state = self.graph.invoke(initial_state)
            
            response = {
                "answer": final_state.get("final_answer", "No answer could be generated."),
                "confidence": final_state.get("confidence", 0.0),
                "selected_groups": final_state.get("selected_groups", []),
                "search_results": final_state.get("search_results", []),
            }
            # Don't cache here, as the full conversational context might be needed next time.
            # self.query_cache[cache_key] = response 
            return response
            
        except Exception as e:
            logger.error(f"Error processing query graph: {e}", exc_info=True)
            return {
                "answer": f"I encountered an error: {str(e)}",
                "confidence": 0.0, "selected_groups": [], "search_results": []
            }

    def _select_groups(self, state: AgentState) -> AgentState:
        question = state["question"]
        conversation_history = self._format_history(state["messages"])
        
        # Fallback to AI group selection if the summary file check fails or yields no results.
        group_selection_prompt = f"""You are an expert document routing system for oil and gas operations. Based on the conversation history and the current question, select the most relevant document groups.

        Conversation History:
        ---
        {conversation_history}
        ---

        Available Groups and Their Precise Content:
        - **Chemical Consumption:** ONLY for data on chemicals, additives, and fluid treatments.
        - **Engineering Tickets:** ONLY for technical issues, equipment failures, and mechanical problems.
        - **Risk Assessment and Hazard Analysis:** ONLY for safety documents, hazard analysis (HAZID, HAZOP), incident reports, and risk control.
        - **Well Recap:** ONLY for well completion summaries, drilling progress, dates, mud cost, and final well reports.
        - **Mud Program:** ONLY for drilling mud properties, fluid engineering, and mud reports.
        - **Hydraulic Summary:** ONLY for data on pressure, flow rates, and pump operations.
        - **Other Group:** A default for general documents that don't fit elsewhere.

        Current User Question: "{question}"

        Instructions:
        1.  **Analyze the User's Intent:** What is the core subject of the current question? Use the history for context.
        2.  **Consider Overlap:** If a question involves multiple domains, select all relevant groups.
        3.  **Provide a Final Selection:** Return ONLY a single JSON array of the chosen group names.

        Selected Groups (JSON Array):"""

        try:
            response = self.azure_client.chat.completions.create(
                model=self.azure_deployment_name,
                messages=[{"role": "user", "content": group_selection_prompt}],
                max_tokens=200,
                temperature=0
            )
            
            groups_text = response.choices[0].message.content.strip()
            json_match = re.search(r'\[.*?\]', groups_text, re.DOTALL)
            if not json_match:
                raise ValueError("LLM did not return a valid JSON array string.")
                
            selected_groups = json.loads(json_match.group(0))
            
            valid_groups = [g for g in selected_groups if g in FILE_GROUPS]
            if not valid_groups:
                logger.warning(f"LLM selected no valid groups. Defaulting to all groups.")
                valid_groups = list(FILE_GROUPS.keys())
            
            logger.info(f"Selected groups for question '{question}' via LLM: {valid_groups}")
            state["selected_groups"] = valid_groups
            
        except Exception as e:
            logger.error(f"Error selecting groups with LLM, defaulting to all. Error: {e}")
            state["selected_groups"] = list(FILE_GROUPS.keys())
            
        return state
    
    # In class LangGraphAgent:
    def _search_documents(self, state: AgentState) -> dict:
        """
        --- HEAVILY MODIFIED FUNCTION ---
        Implements a robust search with a fallback mechanism.
        1. Tries a precise "filter-then-rank" search using dates and entities.
        2. If #1 yields no results and filters were active, it performs a broad 
        vector search as a fallback.
        """
        logger.info("Performing enhanced document search...")
        
        selected_groups = state["selected_groups"]
        start_date = state.get("start_date")
        end_date = state.get("end_date")
        query_entities = state.get("query_entities", {})
        original_question = state["original_question"]
        
        # --- PATH 1: PRECISE, FILTERED SEARCH ---
        logger.info("Attempting precise 'filter-then-rank' search...")
        
        # Get all chunks from the selected groups to create a candidate pool
        candidate_chunks = []
        for group_name in selected_groups:
            if group_name in self.vector_store.group_chunks:
                candidate_chunks.extend(self.vector_store.group_chunks[group_name])
                
        logger.info(f"Initial candidate pool from selected groups: {len(candidate_chunks)} chunks.")

        # Apply date and entity filters
        filters_were_applied = bool(start_date or any(query_entities.values()))
        if filters_were_applied:
            if start_date and end_date:
                candidate_chunks = [c for c in candidate_chunks if any(start_date <= d <= end_date for d in c.metadata.get("dates", []))]
                logger.info(f"After date filter: {len(candidate_chunks)} chunks remain.")
            
            wells_to_match = query_entities.get("wells", [])
            if wells_to_match:
                candidate_chunks = [c for c in candidate_chunks if any(w.lower() in (mw.lower() for mw in c.metadata.get("wells", [])) for w in wells_to_match)]
                logger.info(f"After 'well' entity filter: {len(candidate_chunks)} chunks remain.")

        final_results = []
        if candidate_chunks:
            # Re-rank the filtered candidates
            sentence_pairs = [[original_question, chunk.content] for chunk in candidate_chunks]
            if len(sentence_pairs) > 200:
                logger.warning(f"Clipping candidates for re-ranking from {len(sentence_pairs)} to 200.")
                sentence_pairs = sentence_pairs[:200]
                candidate_chunks = candidate_chunks[:200]

            rerank_scores = self.cross_encoder.predict(sentence_pairs)
            re_ranked_results = sorted(zip(candidate_chunks, rerank_scores), key=lambda x: x[1], reverse=True)
            
            # Format for "Small-to-Big" synthesis
            top_results = re_ranked_results[:15]
            seen_parent_content = set()
            for chunk, score in top_results:
                parent_content = chunk.metadata.get('parent_content', chunk.content)
                if parent_content not in seen_parent_content:
                    final_results.append({
                        "child_content": chunk.content, "content_for_llm": parent_content,
                        "file_name": chunk.metadata.get('file_name', 'unknown'), "group": chunk.group,
                        "score": float(score), "chunk_id": chunk.chunk_id, "metadata": chunk.metadata
                    })
                    seen_parent_content.add(parent_content)

        # --- PATH 2: FALLBACK BROAD SEARCH ---
        # Trigger if the precise search found nothing AND we had active filters
        if not final_results and filters_were_applied:
            logger.warning("Filtered search yielded no results. Performing a broad fallback vector search.")
            
            all_broad_results = []
            seen_chunk_ids = set()
            queries = state.get("sub_queries", [original_question])

            # Run vector search for each sub-query, without metadata filters
            for query in queries:
                # k=20 per query, we will re-rank later
                broad_candidate_results = self.vector_store.search_groups(query, selected_groups, k=20)
                for res_chunk, res_score, res_group in broad_candidate_results:
                    if res_chunk.chunk_id not in seen_chunk_ids:
                        all_broad_results.append((res_chunk, res_score, res_group))
                        seen_chunk_ids.add(res_chunk.chunk_id)

            if not all_broad_results:
                return {"search_results": [], "fallback_triggered": False}

            # Re-rank the new broad candidates
            fallback_sentence_pairs = [[original_question, chunk.content] for chunk, _, _ in all_broad_results]
            fallback_rerank_scores = self.cross_encoder.predict(fallback_sentence_pairs)
            
            fallback_ranked_results = sorted(zip(all_broad_results, fallback_rerank_scores), key=lambda x: x[1], reverse=True)

            # Format for "Small-to-Big" synthesis
            top_fallback_results = fallback_ranked_results[:15]
            seen_parent_content = set()
            for (chunk_tuple, score) in top_fallback_results:
                chunk, _, _ = chunk_tuple # unpack the original result tuple
                parent_content = chunk.metadata.get('parent_content', chunk.content)
                if parent_content not in seen_parent_content:
                    final_results.append({
                        "child_content": chunk.content, "content_for_llm": parent_content,
                        "file_name": chunk.metadata.get('file_name', 'unknown'), "group": chunk.group,
                        "score": float(score), "chunk_id": chunk.chunk_id, "metadata": chunk.metadata
                    })
                    seen_parent_content.add(parent_content)
            
            logger.info(f"Fallback search found and re-ranked {len(final_results)} unique parent contexts.")
            return {"search_results": final_results, "fallback_triggered": True}

        # Return results from the primary search path
        logger.info(f"Primary search found and re-ranked {len(final_results)} unique parent contexts.")
        return {"search_results": final_results, "fallback_triggered": False}

    def _extract_query_date_range(self, state: AgentState) -> dict: # Note: return type hint is now dict
        """
        NEW NODE: Use an LLM to extract start and end dates from the user's query.
        """
        logger.info("Extracting date range from query...")
        question = state["original_question"]

        date_extraction_prompt = f"""
        Analyze the user's question to extract a start date and an end date.
        The current date is {datetime.now().strftime('%Y-%m-%d')}.
        - If you find a date range, return a JSON object with "start_date" and "end_date" in "YYYY-MM-DD" format.
        - If you find only one date, use it for both start and end.
        - If no specific dates or ranges are mentioned, return an empty JSON object {{}}.

        User Question: "{question}"

        JSON Output:
        """
        
        try:
            response = self.azure_client.chat.completions.create(
                model=self.azure_deployment_name,
                messages=[{"role": "user", "content": date_extraction_prompt}],
                temperature=0.0,
                response_format={"type": "json_object"}
            )
            dates = json.loads(response.choices[0].message.content)
            logger.info(f"Extracted date range: {dates}")
            
            # --- START OF FIX ---
            # Return ONLY the new keys. Do not return the whole state object.
            return {
                "start_date": dates.get("start_date"),
                "end_date": dates.get("end_date")
            }
            # --- END OF FIX ---

        except Exception as e:
            logger.error(f"Date range extraction failed: {e}. Skipping date filter.")
            # Return None for the new keys if extraction fails
            return {"start_date": None, "end_date": None}

    def _extract_entities(self, state: AgentState) -> dict:
        """
        --- NEW FUNCTION ---
        Extracts structured entities from the query for metadata-based pre-filtering.
        """
        logger.info("Extracting structured entities from query...")
        question = state["original_question"]
        
        # --- NEW ---
        # Prompt to extract specific, filterable entities.
        entity_extraction_prompt = f"""
        Analyze the user's question to extract key filterable entities.
        Extract any specific well names (e.g., "Well-A123", "G-45"), or specific chemical names or dates, costs
        Return a single JSON object with the keys "wells" and "chemicals".
        If no entities of a certain type are found, return an empty list for that key.

        User Question: "{question}"

        Example for "What were the safety issues with barite on Well-A123?":
        {{
            "wells": ["Well-A123"],
            "chemicals": ["barite"]
            "dates" : ["21/07/2013"]   
            "costs" : ["$1000"]   }}

        JSON Output:
        """
        
        try:
            response = self.azure_client.chat.completions.create(
                model=self.azure_deployment_name,
                messages=[{"role": "user", "content": entity_extraction_prompt}],
                temperature=0.0,
                response_format={"type": "json_object"}
            )
            entities = json.loads(response.choices[0].message.content)
            
            # Ensure the output is a dictionary with lists
            validated_entities = {
                "wells": [w for w in entities.get("wells", []) if isinstance(w, str)],
                "chemicals": [c for c in entities.get("chemicals", []) if isinstance(c, str)]
            }
            logger.info(f"Extracted and validated entities: {validated_entities}")
            return {"query_entities": validated_entities} # Add to state
            
        except (json.JSONDecodeError, KeyError, Exception) as e:
            logger.error(f"Entity extraction failed: {e}. Proceeding without entity filters.")
            return {"query_entities": {}}

    ## NEW: Node for refining the answer based on feedback, using the new format.
    def _refine_answer(self, state: AgentState) -> AgentState:
        # This method's prompt is updated to include history for better refinement.
        logger.info(f"Refining answer (Attempt {state['refinement_attempts']})...")
        
        question = state["original_question"]
        conversation_history = self._format_history(state["messages"])
        context_parts = []
        for res in state["search_results"]:
            context_parts.append(f"Source (File: {res['file_name']}, Chunk ID: {res['chunk_id']}):\n{res['content_for_llm']}\n---")
        context_str = "\n".join(context_parts)
        previous_answer = state["final_answer"]
        feedback = state["critic_feedback"]
        
        refine_prompt = f"""You are an expert Q&A system. You previously generated an answer that was critiqued. 
        Your task is to generate a new, improved answer that addresses the feedback, considering the full conversation history and following the strict format below.

        **Conversation History:**
         ---
          {conversation_history}
        ---

        **Original Question:** "{question}"

        **Provided Context from Documents:**
        ---
        {context_str}
        ---

        **Your Previous Answer:**
        "{previous_answer}"

        **Critique and Instructions for Improvement:**
        "{feedback}"

        **Required Output Format (Strict):**

        [Start with a direct, one-sentence answer, with the key finding in bold. Then, provide additional relevant context, incorporating the feedback. Cite your sources using <cite id="chunk_id_goes_here"></cite>.]

        ---

        ### Citation
         * **File Name**: `[List the name of the file or files the primary information came from.]`
        ### More Information
        * **Supporting Quote**: `[Quote the specific text from the document that directly supports the main answer.]`

        Please provide the new, revised answer in the format above.

        **Revised Answer:**"""

        response = self.azure_client.chat.completions.create(
            model=self.azure_deployment_name,
            messages=[{"role": "user", "content": refine_prompt}],
            max_tokens=1500,
            temperature=0.0
        )
        revised_answer = response.choices[0].message.content
        
        state["final_answer"] = revised_answer
        return state


class AgenticRAGWithLangGraph:
    """Main Agentic RAG system class"""
    
    ## MODIFIED: Accepts azure_deployment_name to be passed to the agent.
    def __init__(self, azure_endpoint: str, azure_key: str, api_version: str, deepseek_api_key: str, azure_deployment_name: str):
        # ... __init__ is unchanged ...
        self.azure_client = AzureOpenAI(
            azure_endpoint=azure_endpoint,
            api_key=azure_key,
            api_version=api_version
        )
        self.deepseek_client = OpenAI(
            api_key=deepseek_api_key,
            base_url="https://api.deepseek.com/v1"
        )
        self.azure_deployment_name = azure_deployment_name
        self.document_processor = DocumentProcessor(self.azure_client)
        self.vector_store = GroupedVectorStore()
        self.file_analyzer = None
        self.agent = None
        logger.info("Initialized AgenticRAG with LangGraph system and DeepSeek client")

    def extract_structured_data(self, file_path: str, full_text: str) -> Optional[Dict]:
        """Uses an LLM to extract structured data from a document's text."""
        if not full_text.strip():
            return None

        prompt = f"""You are an automated data extraction engine for the oil and gas industry. Analyze the following document content and extract key information.

        Return a single JSON object with the following fields:
        - "document_type": Classify the document. Options are "Well Recap", "Chemical Consumption", "Engineering Ticket", "Risk Assessment", "Mud Program", "Hydraulic Summary", or "Other".
        - "key_entities": A dictionary of critical entities found. Include "wells" (list of well names), "dates" (list of specific dates), and "chemicals" (list of key chemical names).
        - "summary": A concise, one-sentence summary of the document's core purpose.

        Document Content:
        \"\"\"
        {full_text[:4000]}
        \"\"\"

        JSON Output:
        """
        try:
            response = self.azure_client.chat.completions.create(
                model="gpt-4o",  # Or your Azure deployment name
                messages=[{"role": "user", "content": prompt}],
                temperature=0.0,
                response_format={"type": "json_object"}
            )
            extracted_data = json.loads(response.choices[0].message.content)
            extracted_data['file_name'] = os.path.basename(file_path)  # Add file name for linking
            return extracted_data
        except Exception as e:
            logger.error(f"Structured extraction failed for {file_path}: {e}")
            return None


    def get_system_stats(self) -> Dict[str, Any]:
        """Enhanced system statistics"""
        stats = {
            "vector_store_initialized": bool(self.vector_store.group_indices),
            "total_groups": len(self.vector_store.group_indices),
            "agent_initialized": self.agent is not None,
            "file_analyzer_available": self.file_analyzer is not None,
            "supported_file_types": list(self.document_processor.supported_extensions),
            "embedding_model": self.vector_store.model_name if self.vector_store else "Unknown"
        }
        
        if self.vector_store.group_indices:
            enhanced_stats = self.vector_store.get_enhanced_group_stats()
            
            # --- FIX: Call the correct method on self.vector_store ---
            file_summary = self.vector_store.get_file_type_summary()
            
            stats.update({
                "file_type_summary": file_summary,
                "enhanced_group_statistics": enhanced_stats['groups'],
                "file_type_distribution": enhanced_stats['file_type_distribution'],
                "total_chunks": sum(g['chunk_count'] for g in enhanced_stats['groups'].values()),
                "total_unique_files": file_summary.get('total_files', 0), # Get from file_summary
                "available_groups": list(self.vector_store.group_indices.keys())
            })
            
            if hasattr(self.document_processor, 'processing_stats'):
                stats["processing_statistics"] = self.document_processor.processing_stats
        
        return stats
    
    def ingest_folder(self, folder_path: str, save_path: str = "./grouped_vector_store"):
        """Ingest all documents from a folder with group classification and structured data extraction."""
        logger.info(f"Starting grouped ingestion of folder: {folder_path}")

        if not os.path.exists(folder_path):
            logger.error(f"Folder does not exist: {folder_path}")
            return 0

        # Create or clear the summary index file
        summary_index_path = "summary_index.jsonl"
        with open(summary_index_path, 'w') as f:
            pass  # Creates an empty file to start

        self.file_analyzer = FileSystemAnalyzer(folder_path)
        self.file_analyzer.set_document_processor(self.document_processor)
        self.file_analyzer.analyze()

        all_chunks = []
        processed_files = 0
        group_stats = {group: 0 for group in FILE_GROUPS.keys()}

        for root, dirs, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                file_ext = Path(file).suffix.lower()
                chunks = []
                try:
                    if file_ext == '.pdf':
                        chunks = self.document_processor.process_pdf(file_path)
                    elif file_ext == '.docx':
                        chunks = self.document_processor.process_docx(file_path)
                    elif file_ext == '.txt':
                        chunks = self.document_processor.process_txt(file_path)
                    else:
                        continue

                    if chunks:
                        all_chunks.extend(chunks)
                        processed_files += 1
                        for chunk in chunks:
                            group_stats[chunk.group] += 1
                        
                        # Extract structured data from the full text of the document
                        full_text = " ".join([c.content for c in chunks])
                        structured_data = self.extract_structured_data(file_path, full_text)

                        if structured_data:
                            with open(summary_index_path, 'a') as f:
                                f.write(json.dumps(structured_data) + '\n')
                        
                        logger.info(f"Processed {file_path}: {len(chunks)} chunks in group {chunks[0].group}")
                    else:
                        logger.warning(f"No chunks created from {file_path}")

                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
                    continue

        if not all_chunks:
            logger.error("No chunks were created from any files")
            return 0

        self.vector_store.add_documents(all_chunks)
        self.vector_store.save_all_groups(save_path)

        self.agent = LangGraphAgent(
            azure_client=self.azure_client,
            deepseek_client=self.deepseek_client,
            vector_store=self.vector_store,
            file_analyzer=self.file_analyzer,
            azure_deployment_name=self.azure_deployment_name
        )

        logger.info(f"Ingestion complete. Processed {processed_files} files, created {len(all_chunks)} chunks.")
        logger.info(f"Group distribution: {group_stats}")
        return len(all_chunks)
    
    def load_vector_store(self, save_path: str = "./grouped_vector_store") -> Dict[str, Any]:
        """Load a previously saved grouped vector store and return detailed results"""
        logger.info(f"Loading grouped vector store from: {save_path}")
        result = {
            "success": False, "loaded_groups": [], "error_message": None, "groups_count": 0
        }
        try:
            loaded_groups = self.vector_store.load_all_groups(save_path)
            if loaded_groups:
                # Create a minimal file analyzer if one doesn't exist
                if not self.file_analyzer:
                    self.file_analyzer = FileSystemAnalyzer("./")
                    self.file_analyzer.set_document_processor(self.document_processor)

                # ## MODIFIED: Pass the deepseek_client to the agent
                self.agent = LangGraphAgent(
                azure_client=self.azure_client,
                deepseek_client=self.deepseek_client,
                vector_store=self.vector_store,
                file_analyzer=self.file_analyzer,
                azure_deployment_name=self.azure_deployment_name # ## MODIFIED: Pass deployment name
            )

                result.update({
                    "success": True, "loaded_groups": loaded_groups, "groups_count": len(loaded_groups)
                })
                logger.info(f"Grouped vector store loaded successfully. Groups: {loaded_groups}")
            else:
                result["error_message"] = "No groups found in the specified path"
                logger.error("Failed to load grouped vector store: No groups found")
        except Exception as e:
            result["error_message"] = str(e)
            logger.error(f"Error loading grouped vector store: {e}")
        return result


    def query(self, messages: List[Dict[str, str]]) -> Dict[str, Any]:
        """Process a query using the LangGraph agentic approach with chat history."""
        if not messages:
            return { "answer": "Please ask a question.", "confidence": 0.0, "selected_groups": [], "search_results": [] }

        logger.info(f"Processing query with history. Last message: {messages[-1]['content']}")
        
        if not self.vector_store.group_indices or not self.agent:
            return {
                "answer": "System not initialized. Please ingest documents first.",
                "confidence": 0.0, "selected_groups": [], "search_results": []
            }

        # Convert the session state messages (dicts) to LangChain message objects
        langchain_messages = []
        for msg in messages:
            if msg["role"] == "user":
                langchain_messages.append(HumanMessage(content=msg["content"]))
            elif msg["role"] == "assistant":
                langchain_messages.append(AIMessage(content=msg["content"]))

        try:
            # The agent's process_query now expects a list of BaseMessage objects
            response = self.agent.process_query(langchain_messages)
            return response
        except Exception as e:
            logger.error(f"Error processing query: {e}", exc_info=True)
            return {
                "answer": f"I apologize, but I encountered an error: {str(e)}",
                "confidence": 0.0, "selected_groups": [], "search_results": []
            }

    
    def get_group_file_details(self, group_name: str) -> Dict[str, Any]:
        """Get detailed file information for a specific group"""
        return self.vector_store.get_group_files(group_name)

    def batch_query(self, questions: List[str]) -> List[Dict[str, Any]]:
        """Process multiple queries in batch"""
        results = []
        
        for i, question in enumerate(questions):
            logger.info(f"Processing batch query {i+1}/{len(questions)}: {question}")
            try:
                response = self.query(question)
                results.append({
                    "question": question,
                    "answer": response["answer"],
                    "confidence": response["confidence"],
                    "selected_groups": response["selected_groups"],
                    "status": "success"
                })
            except Exception as e:
                logger.error(f"Error in batch query {i+1}: {e}")
                results.append({
                    "question": question,
                    "answer": f"Error: {str(e)}",
                    "confidence": 0.0,
                    "selected_groups": [],
                    "status": "error"
                })
        
        return results
    
    def export_knowledge_base(self, export_path: str):
        """Export the grouped knowledge base to a readable format"""
        try:
            export_data = {
                "metadata": {
                    "export_timestamp": datetime.now().isoformat(),
                    "total_groups": len(self.vector_store.group_indices),
                    "system_stats": self.get_system_stats()
                },
                "groups": {}
            }
            
            # Export each group's documents
            for group_name, chunks in self.vector_store.group_chunks.items():
                if not chunks:
                    continue
                
                # Group chunks by file
                files_dict = {}
                for chunk in chunks:
                    file_name = chunk.metadata.get('file_name', 'unknown')
                    if file_name not in files_dict:
                        files_dict[file_name] = []
                    files_dict[file_name].append(chunk)
                
                # Export each file's content
                group_data = {
                    "group_name": group_name,
                    "total_chunks": len(chunks),
                    "total_files": len(files_dict),
                    "documents": []
                }
                
                for file_name, file_chunks in files_dict.items():
                    file_data = {
                        "file_name": file_name,
                        "total_chunks": len(file_chunks),
                        "file_type": file_chunks[0].metadata.get('file_type', 'unknown'),
                        "processed_at": file_chunks[0].metadata.get('processed_at', 'unknown'),
                        "chunks": [
                            {
                                "chunk_id": chunk.chunk_id,
                                "content": chunk.content,
                                "metadata": chunk.metadata
                            } for chunk in file_chunks
                        ]
                    }
                    group_data["documents"].append(file_data)
                
                export_data["groups"][group_name] = group_data
            
            # Save to JSON
            with open(export_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Grouped knowledge base exported to {export_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error exporting knowledge base: {e}")
            return False

    def save_vector_store(self, save_path: str) -> bool:
        """Saves all groups from the vector store to a folder."""
        if not self.vector_store:
            logger.error("Vector store not initialized, cannot save.")
            return False
        return self.vector_store.save_all_groups(save_path)

def enhanced_sidebar_display(st, rag_system):
    """Enhanced sidebar with professional styling"""
    
    # Custom sidebar styling
    st.markdown("""
    <style>
    .sidebar .sidebar-content {
        background-color: #2c3e50;
        color: white;
    }
    .sidebar .sidebar-content .stRadio > div > label {
        color: white;
    }
    .sidebar .sidebar-content .stTextInput > div > div > input {
        background-color: #34495e;
        color: white;
    }
    .sidebar .sidebar-content .stButton > button {
        background-color: #3498db;
        color: white;
        border: none;
        border-radius: 4px;
        padding: 8px 16px;
    }
    .sidebar .sidebar-content .stButton > button:hover {
        background-color: #2980b9;
    }
    .sidebar .sidebar-content .stExpander > div > div {
        background-color: #34495e;
        border-radius: 4px;
        padding: 8px;
        margin-bottom: 8px;
    }
    </style>
    """, unsafe_allow_html=True)
    
    with st.sidebar:
        # Sidebar header with accent color
        st.markdown("""
        <div style="background-color: #3498db; padding: 10px; border-radius: 4px; margin-bottom: 20px;">
            <h2 style="color: white; margin: 0;">Configuration</h2>
        </div>
        """, unsafe_allow_html=True)
        
        # Action selection with custom styling
        action = st.radio("Choose Action", ["Ingest New Folder", "Load Vector Store"])
        
        if action == "Ingest New Folder":
            folder_path = st.text_input("Enter folder path to ingest:", "./Input")
            if st.button("Ingest Documents", key="ingest_btn") and folder_path:
                if os.path.exists(folder_path):
                    with st.spinner(f"Ingesting documents from {folder_path}..."):
                        chunks_created = rag_system.ingest_folder(folder_path)
                        st.success(f"Created {chunks_created} chunks from documents")
                        st.rerun()
                else:
                    st.error(f"Folder {folder_path} does not exist.")
        
        elif action == "Load Vector Store":
            # CORRECTED LINE
            vector_store_path = st.text_input("Enter vector store path:", "./grouped_vector_store")
            if st.button("Load Vector Store", key="load_btn"):
                with st.spinner("Loading vector store..."):
                    if rag_system.load_vector_store(vector_store_path):
                        st.success("Vector store loaded successfully")
                        st.rerun()
                    else:
                        st.error("Failed to load vector store")
        
        # Divider with custom style
        st.markdown("""
        <div style="height: 1px; background-color: #3498db; margin: 20px 0;"></div>
        """, unsafe_allow_html=True)
        
        # Save section with custom styling
        st.markdown("""
        <div style="background-color: #34495e; padding: 10px; border-radius: 4px; margin-bottom: 20px;">
            <h3 style="color: white; margin: 0;">Save Vector Store</h3>
        </div>
        """, unsafe_allow_html=True)
        
        save_path = st.text_input("Save path:", "./grouped_vector_store")
        if st.button("Save Vector Store", key="save_btn"):
            if rag_system.save_vector_store(save_path):
                st.success("Vector store saved successfully")
            else:
                st.error("Failed to save vector store")
        
        # Document groups section
        st.markdown("""
        <div style="background-color: #34495e; padding: 10px; border-radius: 4px; margin: 20px 0;">
            <h3 style="color: white; margin: 0;">Document Groups</h3>
        </div>
        """, unsafe_allow_html=True)
        
        stats = rag_system.get_system_stats()
        
        for group_name, keywords in FILE_GROUPS.items():
            group_stats = stats.get('enhanced_group_statistics', {}).get(group_name, {})
            file_count = group_stats.get('unique_files', 0)
            
            with st.expander(f"{group_name} ({file_count} files)"):
                if keywords:
                    st.markdown(f"<p style='color: #3498db;'><strong>Keywords:</strong> {', '.join(keywords)}</p>", unsafe_allow_html=True)
                else:
                    st.markdown("<p style='color: #3498db;'><strong>Default group</strong> for unmatched files</p>", unsafe_allow_html=True)
                
                if file_count > 0:
                    st.markdown(f"<p><strong>Total chunks:</strong> {group_stats.get('chunk_count', 0)}</p>", unsafe_allow_html=True)
                    
                    file_details = rag_system.get_group_file_details(group_name)
                    
                    st.markdown("<p><strong>Files in this group:</strong></p>", unsafe_allow_html=True)
                    for file_name in file_details.get('files', []):
                        file_info = file_details.get('file_details', {}).get(file_name, {})
                        processing_method = file_info.get('processing_method', 'Unknown')
                        chunks = file_info.get('chunks', 0)
                        
                        st.markdown(f"<p style='margin-left: 20px;'>{file_name} ({chunks} chunks)</p>", unsafe_allow_html=True)
                        
                        if file_info.get('total_pages', 0) > 0:
                            vision_pages = len(file_info.get('vision_pages', []))
                            text_pages = len(file_info.get('text_pages', []))
                            st.markdown(f"<p style='margin-left: 20px;'>Pages: {text_pages} text, {vision_pages} OCR</p>", unsafe_allow_html=True)
                
                else:
                    st.markdown("<p style='color: #7f8c8d;'>No files in this group</p>", unsafe_allow_html=True)

## MODIFIED: Re-integrated the statistics dashboard into the chat interface.
def enhanced_main_display(st, rag_system):
    """Enhanced main display with a statistics dashboard and a conversational chat interface."""
    
    # Custom main page styling (unchanged)
    st.markdown("""
    <style>
    .main .block-container {
        padding-top: 2rem;
    }
    .stMetric {
        background-color: white;
        border-radius: 8px;
        padding: 15px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    /* Add other styles from previous versions if needed */
    </style>
    """, unsafe_allow_html=True)

    # Main header with gradient background
    st.markdown("""
    <div style="background: linear-gradient(135deg, #3498db, #2c3e50); padding: 20px; border-radius: 8px; margin-bottom: 30px; color: white;">
        <h1 style="color: white; margin: 0;">Agentic RAG Chat 🤖</h1>
        <p style="margin: 5px 0 0 0;">Ask questions about your documents</p>
    </div>
    """, unsafe_allow_html=True)

    # --- NEW: Re-added Statistics Display ---
    stats = rag_system.get_system_stats()
    
    # Display file type and processing method breakdown if a vector store is loaded
    if stats.get('vector_store_initialized') and stats.get('file_type_summary'):
        st.markdown("### 📊 System Dashboard")
        file_summary = stats['file_type_summary']
        
        # Display File Processing Summary
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("📄 PDF Files Processed", file_summary.get('pdf_files', 0))
        with col2:
            st.metric("📃 DOCX Files Processed", file_summary.get('docx_files', 0))
        with col3:
            st.metric("📝 TXT Files Processed", file_summary.get('txt_files', 0))

        # Display Processing Methods Summary
        proc_methods = file_summary.get('processing_methods', {})
        col4, col5 = st.columns(2)
        with col4:
            st.metric("🔤 Text Extraction", proc_methods.get('text_extraction', 0), help="Files/pages processed using standard text extraction.")
        with col5:
            st.metric("👁️ Vision (OCR)", proc_methods.get('vision_ocr', 0), help="Files/pages processed using vision-based OCR for scanned content.")
        
        st.markdown("---") # Visual separator

    # --- Chat Interface (from previous update) ---
    st.markdown("### 💬 Chat")
    
    # Display chat messages from history
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"], unsafe_allow_html=True)

    # Accept user input
    if prompt := st.chat_input("Ask a follow-up or a new question..."):
        # Add user message to chat history
        st.session_state.messages.append({"role": "user", "content": prompt})
        # Display user message
        with st.chat_message("user"):
            st.markdown(prompt)

        # Get and display assistant response
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                response = rag_system.query(st.session_state.messages)
                
                answer_html = response['answer']
                
                # Parse for <cite> tags and create a citation mapping
                cite_matches = list(re.finditer(r'<cite id="(.*?)"></cite>', answer_html))
                citations = {}
                for i, match in enumerate(cite_matches):
                    chunk_id = match.group(1)
                    cite_key = f"[{i+1}]"
                    citations[cite_key] = chunk_id
                    # Replace the tag with a superscript number
                    answer_html = answer_html.replace(match.group(0), f"&nbsp;<sup>**{cite_key}**</sup>")

                # Display the main answer with citation links
                st.markdown(answer_html, unsafe_allow_html=True)

                # Display evidence in an expander
                if response.get('search_results'):
                    with st.expander("View Evidence and Citations"):
                        if citations:
                            st.markdown("#### Citations")
                            for key, chunk_id in citations.items():
                                st.markdown(f"**{key}**: `{chunk_id}`")
                        
                        st.markdown("---")
                        st.markdown("#### Top Retrieved Documents")
                        for result in response['search_results'][:3]:
                            st.markdown(f"**File:** `{result['file_name']}` (Score: {result['score']:.3f})")
                            st.text_area(
                                label=f"Content from Chunk ID: {result['chunk_id']}",
                                value=result.get('content_for_llm', result.get('child_content', 'Content not available.')),
                                height=150,
                                disabled=True,
                                key=f"evidence_{result['chunk_id']}"
                            )
                
                # Add the complete assistant response (with HTML for citations) to history
                st.session_state.messages.append({"role": "assistant", "content": answer_html})



def main():
    """Main function to run the enhanced Streamlit demo with chat."""
    st.set_page_config(page_title="Advanced Agentic RAG", layout="wide")

    # --- API Key Configuration (Unchanged) ---
    AZURE_ENDPOINT = "https://oai-nasco.openai.azure.com/"
    AZURE_KEY = st.secrets.get("AZURE_API_KEY") 
    API_VERSION = "2024-02-15-preview"
    AZURE_DEPLOYMENT_NAME = "gpt-4o"
    DEEPSEEK_API_KEY = st.secrets.get("DEEPSEEK_API_KEY")

    if not DEEPSEEK_API_KEY:
        st.error("DeepSeek API key is missing! Please add it to your secrets.")
        st.stop()
    
    # Initialize the RAG system in the session state if it's not already there
    if 'rag_system' not in st.session_state:
        st.session_state.rag_system = AgenticRAGWithLangGraph(
            azure_endpoint=AZURE_ENDPOINT, 
            azure_key=AZURE_KEY, 
            api_version=API_VERSION, 
            deepseek_api_key=DEEPSEEK_API_KEY,
            azure_deployment_name=AZURE_DEPLOYMENT_NAME
        )
    
    # ## NEW: Initialize chat history in session state
    if 'messages' not in st.session_state:
        st.session_state.messages = []
        
    rag_system = st.session_state.rag_system
    
    # Display the sidebar (unchanged)
    enhanced_sidebar_display(st, rag_system)
    
    # Display the main chat interface
    # This function now reads/writes to st.session_state.messages directly
    enhanced_main_display(st, rag_system)


if __name__ == "__main__":
    main()
